"""
TwinBros_04_Guia_Completa.docx
Documento único: todo lo que tiene la web + mejoras posibles + tarifas estándar de web
Para uso interno del equipo comercial
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = r"C:\Users\jose2\OneDrive\Escritorio\mcp\TwinBros_04_Guia_Completa_y_Tarifas.docx"

# ── Paleta Twin Bros ──────────────────────────────────────────────────────────
MAGENTA  = RGBColor(224, 64, 251)
NEGRO    = RGBColor(10, 10, 10)
GRIS     = RGBColor(100, 100, 100)
BLANCO   = RGBColor(255, 255, 255)
VERDE    = RGBColor(30, 120, 60)
AZUL     = RGBColor(26, 58, 107)
ROJO     = RGBColor(180, 30, 30)
DORADO   = RGBColor(200, 146, 10)

def set_margins(doc, top=2, bottom=2, left=2.5, right=2.5):
    s = doc.sections[0]
    s.top_margin = Cm(top); s.bottom_margin = Cm(bottom)
    s.left_margin = Cm(left); s.right_margin = Cm(right)

def cell_shading(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def h1(doc, text, color=None):
    p = doc.add_heading(text, level=1)
    if p.runs:
        r = p.runs[0]
        r.font.color.rgb = color or MAGENTA
    return p

def h2(doc, text, color=None):
    p = doc.add_heading(text, level=2)
    if p.runs:
        r = p.runs[0]
        r.font.color.rgb = color or NEGRO
    return p

def h3(doc, text):
    p = doc.add_heading(text, level=3)
    if p.runs:
        p.runs[0].font.color.rgb = GRIS
    return p

def para(doc, text, bold=False, italic=False, size=11, color=None, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p

def bullet(doc, text, indent=0, bold=False, color=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(indent)
    r = p.runs[0] if p.runs else p.add_run(text)
    if not p.runs:
        pass
    else:
        p.clear()
        r = p.add_run(text)
    r.bold = bold
    if color:
        r.font.color.rgb = color
    return p

def nbullet(doc, text, indent=0.3, bold=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(f'• {text}')
    r.font.size = Pt(10.5)
    r.bold = bold
    if color:
        r.font.color.rgb = color
    return p

def table(doc, headers, rows, header_hex='1A1A1A'):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        c = hrow.cells[i]
        c.text = h
        if c.paragraphs[0].runs:
            run = c.paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = BLANCO
            run.font.size = Pt(10)
        cell_shading(c, header_hex)
    for ri, rd in enumerate(rows):
        row = t.rows[ri + 1]
        for ci, ct in enumerate(rd):
            cell = row.cells[ci]
            cell.text = str(ct)
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(10)
    return t

def highlight_box(doc, text, bg_hex='FFF8E1', color=None):
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    c = t.rows[0].cells[0]
    c.text = text
    if c.paragraphs[0].runs:
        r = c.paragraphs[0].runs[0]
        r.font.size = Pt(10.5)
        r.bold = True
        if color:
            r.font.color.rgb = color
    cell_shading(c, bg_hex)
    doc.add_paragraph()

def divider(doc):
    doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════

def build():
    doc = Document()
    set_margins(doc)

    # ── PORTADA ──────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('TWIN BROS TATTOO STUDIO')
    r.bold = True; r.font.size = Pt(30); r.font.color.rgb = MAGENTA

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('GUIA COMPLETA — WEB + MEJORAS + TARIFAS')
    r2.font.size = Pt(15); r2.font.color.rgb = GRIS

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run('Uso interno · Tomares, Sevilla · Marzo 2026')
    r3.font.size = Pt(10); r3.font.color.rgb = GRIS; r3.italic = True

    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # SECCIÓN 1 — INVENTARIO COMPLETO DE LA WEB ACTUAL
    # ═════════════════════════════════════════════════════════════════════════
    h1(doc, '1. TODO LO QUE TIENE LA WEB AHORA MISMO')
    para(doc, 'Inventario exhaustivo de todo lo entregado a 21/03/2026. URL de produccion: https://twinbros-web.vercel.app', italic=True, size=10, color=GRIS)
    divider(doc)

    # 1.1 ESTRUCTURA
    h2(doc, '1.1 Estructura y paginas')
    table(doc,
        ['Pagina / Ruta', 'Tipo', 'Contenido'],
        [
            ['/ (Inicio)', 'Landing principal', 'Hero + About + Artistas + Galeria + Citas + Contacto'],
            ['/artista/pedro-martin-art', 'Perfil individual', 'Bio extendida + galeria + videos de Pedro'],
            ['/artista/adri-xegg', 'Perfil individual', 'Bio extendida + galeria + videos de Adri'],
            ['/artista/maranhastattoo', 'Perfil individual', 'Bio extendida + galeria + videos de Maranhas'],
            ['/artista/dual-piercing', 'Perfil individual', 'Bio extendida + galeria de Dual Piercing'],
            ['/admin', 'Panel de administracion', 'Gestion de citas, artistas y configuracion'],
            ['/api/content (GET)', 'API REST interna', 'Devuelve el JSON con todo el contenido'],
            ['/api/content (POST)', 'API REST interna', 'Guarda cambios desde el admin'],
        ]
    )
    divider(doc)

    # 1.2 SECCIONES HOME
    h2(doc, '1.2 Secciones de la pagina principal')
    sections_data = [
        ('HERO', 'Portada con video/imagen de fondo, nombre del estudio, boton de reserva, accesos rapidos a los 4 artistas con sus colores identificativos. Animacion de entrada con Framer Motion.'),
        ('ABOUT (Sobre nosotros)', 'Historia de los hermanos fundadores Pedro y Adri, filosofia del estudio, logo de sponsor Big Wasp (@bigwasp.official), valores de marca (Creador + Rebelde).'),
        ('ARTISTAS', 'Grid de 4 tarjetas interactivas con hover. Cada una muestra: foto de perfil, nombre, alias, rol, especialidades con colores, y boton de acceso al perfil completo.'),
        ('GALERIA', 'Filtro por artista (tabs). Fotos y videos reales de Instagram descargados. 17 fotos + 8 videos (2 por artista). Masonry grid responsive.'),
        ('CITAS / RESERVAS', 'Formulario completo: nombre, email, telefono, artista elegido, fecha, descripcion del tatuaje, zona del cuerpo, tamano. Envio al panel admin en tiempo real.'),
        ('CONTACTO', 'Direccion: Alameda Santa Eufemia N32, Tomares. Mapa de Google Maps integrado. Horario: Lu-Vi 10:00-14:00 / 16:00-20:00. Links a Instagram.'),
    ]
    for name, desc in sections_data:
        para(doc, f'▸ {name}', bold=True, size=11)
        para(doc, desc, size=10, indent=0.5)
        doc.add_paragraph()

    # 1.3 PANEL ADMIN
    h2(doc, '1.3 Panel de administracion (/admin)')
    para(doc, 'Acceso con contrasena: twinbros2026', bold=True, size=10)
    doc.add_paragraph()
    table(doc,
        ['Modulo', 'Funciones disponibles'],
        [
            ['Dashboard', 'Contador de citas por estado (pendiente/confirmada/cancelada)'],
            ['Gestion de citas', 'Ver todas las solicitudes, cambiar estado, eliminar, filtrar por artista o estado'],
            ['Gestion de artistas', 'Editar bio, bioExtendida, especialidades, color, redes sociales, foto de perfil'],
            ['Galeria por artista', 'Ver que imagenes y videos tiene cada artista cargados'],
            ['Configuracion', 'Nombre estudio, direccion, horarios, redes sociales, datos del sponsor'],
        ]
    )
    divider(doc)

    # 1.4 CONTENIDO MULTIMEDIA
    h2(doc, '1.4 Contenido multimedia incluido')
    table(doc,
        ['Artista', 'Fotos', 'Videos', 'Extras'],
        [
            ['Pedro Martin Art', '5 trabajos + 2 convenciones (Sevilla + Jerez)', '2 reels', 'Foto perfil, foto convention'],
            ['Adri Xegg', '3 trabajos + 1 brothers + 1 sponsor + 2 convenciones', '2 reels', 'Foto perfil'],
            ['Maranhas', '3 trabajos', '2 reels', 'Foto perfil'],
            ['Dual Piercing', '2 trabajos', 'Sin videos', 'Logo como foto perfil'],
            ['Galeria estudio', '4 fotos de convenciones (Sevilla + Jerez)', '--', 'Fotos generales del estudio'],
        ]
    )
    divider(doc)

    # 1.5 DATOS DE CADA ARTISTA
    h2(doc, '1.5 Datos de cada artista en el sistema')
    table(doc,
        ['Campo', 'Pedro', 'Adri', 'Maranhas', 'Dual Piercing'],
        [
            ['Especialidades', 'Realismo, Black & Grey, Retratos, Gran formato', 'Anime, Neo-trad, Color, Manga, Pop culture', 'Tatuaje artistico, Versatilidad, Int.', 'Piercings, Joyeria corporal, Expansiones'],
            ['Color ID web', 'Azul (#4FC3F7)', 'Naranja (#FF5722)', 'Verde (#66BB6A)', 'Dorado (#FFD700)'],
            ['Instagram', '@pedromartinart', '@adrixegg', '@maranhastattoo', '@dualpiercing'],
            ['TikTok', '@pedromartinart', '@adrixegg', 'No', 'No'],
            ['Seguidores IG', '16K', 'No especificado', 'No especificado', 'No especificado'],
            ['Destacado en web', 'Si', 'Si', 'Si', 'Si'],
        ]
    )
    divider(doc)

    # 1.6 TECNOLOGIA
    h2(doc, '1.6 Stack tecnico')
    table(doc,
        ['Tecnologia', 'Version', 'Para que'],
        [
            ['Next.js', '16', 'Framework React, SSR, rutas dinamicas por artista'],
            ['React', '19', 'UI components'],
            ['TypeScript', '5.x', 'Tipado estatico'],
            ['Tailwind CSS', 'v4', 'Diseño responsive'],
            ['Framer Motion', 'latest', 'Animaciones fluidas'],
            ['Vercel', 'Free plan', 'Hosting, CDN global, deploy automatico'],
            ['JSON file store', '--', 'Base de datos simple en /data/content.json'],
        ]
    )

    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # SECCIÓN 2 — MEJORAS POSIBLES
    # ═════════════════════════════════════════════════════════════════════════
    h1(doc, '2. MEJORAS POSIBLES — CATEGORIAS Y PRECIOS')
    para(doc, 'Clasificadas por impacto, dificultad y precio. Los precios son orientativos segun nuestras tarifas habituales.', italic=True, size=10, color=GRIS)
    divider(doc)

    # CATEGORIA A — Sin coste / inmediatas
    h2(doc, 'CATEGORIA A — Sin coste adicional (incluidas en el mantenimiento basico)')
    highlight_box(doc, 'Estas mejoras se pueden hacer en horas, sin coste extra si hay contrato de mantenimiento.', 'E8F5E9', VERDE)

    mejoras_a = [
        ('Cambiar la contrasena del admin', 'Seguridad basica. 5 minutos.', '0 EUR', '< 1h'),
        ('Actualizar horarios o direccion', 'Si cambian algo del estudio.', '0 EUR', '< 1h'),
        ('Anadir/editar bio de artista', 'Desde el propio admin, sin tocar codigo.', '0 EUR', 'Admin'),
        ('Cambiar foto de perfil de artista', 'Desde el admin del estudio.', '0 EUR', 'Admin'),
        ('Actualizar el sponsor', 'Cambiar logo/link de Big Wasp o poner otro.', '0 EUR', '< 1h'),
        ('Anadir fotos nuevas a la galeria', 'Nuevos trabajos de los artistas.', '0 EUR*', 'Admin'),
    ]
    table(doc, ['Mejora', 'Descripcion', 'Coste', 'Tiempo'], mejoras_a)
    para(doc, '* Las fotos se suben desde el admin pero hay limite de tamano en Vercel (plan gratuito).', italic=True, size=9, color=GRIS)
    divider(doc)

    # CATEGORIA B — Pequenas mejoras
    h2(doc, 'CATEGORIA B — Mejoras pequenas (50–200 EUR)')

    mejoras_b = [
        ('Boton WhatsApp flotante', 'WhatsApp Business directo al numero del estudio. Muchos clientes lo prefieren al formulario.', '60–80 EUR', '2-3h'),
        ('Notificacion por email al recibir cita', 'Email automatico al estudio cuando llega solicitud nueva. Con Resend o EmailJS.', '80–120 EUR', '3-4h'),
        ('Email de confirmacion al cliente', 'El cliente recibe email confirmando que se recibio su solicitud.', '50 EUR (junto al anterior)', '2h extra'),
        ('Dominio personalizado', 'Contratar twinbrostattoo.com o similar. Solo gestion del DNS.', '10-15 EUR/ano + 30 EUR config', '2-3h'),
        ('Nuevo artista (5o artista)', 'Anadir un 5o profesional al estudio con su perfil completo.', '120–180 EUR', '4-6h'),
        ('Formulario de contacto por email', 'Alternativa al formulario de citas para consultas generales.', '80 EUR', '3h'),
        ('Pagina legal (privacidad + cookies)', 'Si quieren cumplir con la RGPD. Cookie banner + politica de privacidad.', '100–150 EUR', '3-4h'),
        ('Google Analytics integrado', 'Ver cuantas visitas tiene la web, de donde vienen, que paginas ven mas.', '60 EUR', '2h'),
        ('Open Graph / redes sociales', 'Imagen previa al compartir en WhatsApp o Instagram. Mejora la imagen de marca.', '50 EUR', '2h'),
    ]
    table(doc, ['Mejora', 'Descripcion', 'Precio', 'Tiempo'], mejoras_b, '2C3E50')
    divider(doc)

    # CATEGORIA C — Mejoras medias
    h2(doc, 'CATEGORIA C — Mejoras medias (200–600 EUR)')

    mejoras_c = [
        ('Galeria con lightbox', 'Las fotos se abren en pantalla completa al hacer clic. Experiencia premium.', '150–200 EUR', '6-8h'),
        ('Filtro por estilo en galeria', 'Ademas de filtrar por artista, filtrar por: Realismo, Anime, Color, B&G, Piercing.', '150–200 EUR', '6h'),
        ('Blog / noticias del estudio', 'Seccion de articulos: convenciones, nuevos artistas, noticias. Bueno para el SEO.', '200–350 EUR', '1-2 dias'),
        ('Panel admin con subida de fotos real', 'Ahora las fotos van por URL. Mejorar a subida directa de archivo (tipo Instagram).', '250–350 EUR', '1 dia'),
        ('Video hero autoplay con reel del estudio', 'El hero carga directamente un video del estudio en lugar de imagen estatica.', '80–120 EUR', '3-4h'),
        ('Calendario de disponibilidad por artista', 'El cliente ve los dias disponibles antes de pedir cita. Reduce citas no viables.', '400–600 EUR', '2-3 dias'),
        ('Multiidioma (ingles/portugues)', 'La web en 2 idiomas. Importante si tienen clientes extranjeros (turismo Sevilla).', '300–500 EUR', '2 dias'),
        ('Testimonios / resenas en la web', 'Seccion con valoraciones de clientes. Puede conectar con Google Reviews.', '150–250 EUR', '6-8h'),
    ]
    table(doc, ['Mejora', 'Descripcion', 'Precio', 'Tiempo'], mejoras_c, '1A3A6B')
    divider(doc)

    # CATEGORIA D — Mejoras grandes
    h2(doc, 'CATEGORIA D — Mejoras grandes (600–2000 EUR)')

    mejoras_d = [
        ('Sistema de citas con deposito online', 'El cliente paga 20-50 EUR de sena al reservar. Elimina no-shows. Stripe integrado.', '600–900 EUR', '3-5 dias'),
        ('Sistema de citas con calendario real', 'Cada artista tiene su agenda. El cliente elige dia/hora disponible. Bloqueo automatico.', '800–1.200 EUR', '5-7 dias'),
        ('Tienda online (prints, aftercare, merch)', 'Si quieren vender merchandising, prints de artistas o productos de cuidado.', '600–1.000 EUR', '3-5 dias'),
        ('Area privada de clientes', 'El cliente se registra y ve su historial de citas, fotos de sus tatuajes, progreso.', '800–1.500 EUR', '1-2 sem'),
        ('Streaming / lives de tattooing', 'Sistema para retransmitir en vivo sesiones de tatuaje desde el estudio.', '500–900 EUR', '3-4 dias'),
        ('App movil PWA avanzada', 'La web funciona como app instalable en el movil con notificaciones push.', '400–700 EUR', '3-5 dias'),
    ]
    table(doc, ['Mejora', 'Descripcion', 'Precio', 'Tiempo'], mejoras_d, '4A0060')
    divider(doc)

    # RESUMEN MEJORAS
    h2(doc, 'Resumen de mejoras por prioridad recomendada')
    highlight_box(doc, 'Recomendacion para empezar: B1 (WhatsApp) + B2 (email citas) + B4 (dominio) = ~200 EUR total. Maximo impacto, minima inversion.', 'FFF8E1', DORADO)

    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # SECCIÓN 3 — TARIFAS ESTANDAR
    # ═════════════════════════════════════════════════════════════════════════
    h1(doc, '3. TARIFAS ESTANDAR — WEBS PARA NEGOCIOS LOCALES')
    para(doc, 'Precios de referencia para propuestas comerciales. Incluyen diseno, desarrollo, deploy en Vercel y onboarding basico.', italic=True, size=10, color=GRIS)
    divider(doc)

    # 3.1 PAQUETES WEB
    h2(doc, '3.1 Paquetes de web para negocio local')

    para(doc, 'PACK BASICO — Landing page sencilla', bold=True, size=12, color=AZUL)
    nbullet(doc, 'Diseno unico con identidad visual del negocio (colores, tipografia, logo)')
    nbullet(doc, 'Secciones: Hero, Sobre nosotros, Servicios, Galeria, Contacto')
    nbullet(doc, 'Formulario de contacto basico')
    nbullet(doc, 'Optimizado para movil')
    nbullet(doc, 'Deploy en Vercel (hosting gratuito)')
    nbullet(doc, 'Sin panel de administracion')
    para(doc, 'Precio: 500–700 EUR (pago unico)', bold=True, size=11, color=VERDE)
    para(doc, 'Tiempo de entrega: 3–5 dias laborables', italic=True, size=10, color=GRIS)
    divider(doc)

    para(doc, 'PACK ESTANDAR — Web completa con admin', bold=True, size=12, color=AZUL)
    nbullet(doc, 'Todo lo del Pack Basico, mas:')
    nbullet(doc, 'Panel de administracion con login y contrasena')
    nbullet(doc, 'Gestion de contenido sin tocar codigo (textos, fotos, horarios)')
    nbullet(doc, 'Fotos y videos reales del negocio (descarga de Instagram/redes incluida)')
    nbullet(doc, 'Animaciones y microinteracciones profesionales')
    nbullet(doc, 'Mapas, redes sociales, botones de accion integrados')
    nbullet(doc, 'SEO basico (meta tags, Open Graph)')
    para(doc, 'Precio: 800–1.200 EUR (pago unico)', bold=True, size=11, color=VERDE)
    para(doc, 'Tiempo de entrega: 5–8 dias laborables', italic=True, size=10, color=GRIS)
    divider(doc)

    para(doc, 'PACK PREMIUM — Web con sistema de citas (tipo Twin Bros)', bold=True, size=12, color=MAGENTA)
    nbullet(doc, 'Todo lo del Pack Estandar, mas:')
    nbullet(doc, 'Sistema de citas online (cliente pide cita, llega al admin)')
    nbullet(doc, 'Gestion de multiples profesionales/servicios')
    nbullet(doc, 'Notificaciones por email al negocio y al cliente')
    nbullet(doc, 'Galeria filtrable por profesional o categoria')
    nbullet(doc, 'Perfil individual por cada profesional')
    nbullet(doc, 'Boton de WhatsApp integrado')
    nbullet(doc, 'Google Analytics configurado')
    para(doc, 'Precio: 1.200–1.800 EUR (pago unico)', bold=True, size=11, color=VERDE)
    para(doc, 'Tiempo de entrega: 7–12 dias laborables', italic=True, size=10, color=GRIS)
    divider(doc)

    para(doc, 'PACK PERSONALIZADO — A medida', bold=True, size=12, color=DORADO)
    nbullet(doc, 'Para proyectos especiales: tienda online, sistema de reservas avanzado, multi-sede')
    nbullet(doc, 'Presupuesto previo sin compromiso')
    para(doc, 'Precio: Desde 2.000 EUR (segun proyecto)', bold=True, size=11, color=VERDE)
    divider(doc)

    # 3.2 TABLA RESUMEN
    h2(doc, '3.2 Tabla comparativa de paquetes')
    table(doc,
        ['Incluye', 'Basico', 'Estandar', 'Premium (Twin Bros)'],
        [
            ['Diseno unico', 'Si', 'Si', 'Si'],
            ['Responsive movil', 'Si', 'Si', 'Si'],
            ['Hosting Vercel', 'Si', 'Si', 'Si'],
            ['Panel admin', 'No', 'Si', 'Si'],
            ['Fotos reales de RRSS', 'No', 'Si', 'Si'],
            ['Animaciones Framer Motion', 'No', 'Si', 'Si'],
            ['Sistema de citas online', 'No', 'No', 'Si'],
            ['Perfiles individuales', 'No', 'No', 'Si'],
            ['Notif. email citas', 'No', 'No', 'Si'],
            ['WhatsApp flotante', 'No', 'No', 'Si'],
            ['SEO basico', 'Basico', 'Estandar', 'Completo'],
            ['PRECIO', '500–700 EUR', '800–1.200 EUR', '1.200–1.800 EUR'],
            ['ENTREGA', '3–5 dias', '5–8 dias', '7–12 dias'],
        ]
    )
    divider(doc)

    # 3.3 SERVICIOS ADICIONALES
    h2(doc, '3.3 Servicios adicionales y mantenimiento')
    table(doc,
        ['Servicio', 'Descripcion', 'Precio'],
        [
            ['Mantenimiento basico mensual', 'Actualizaciones de contenido, cambios de texto/fotos, soporte', '40–60 EUR/mes'],
            ['Mantenimiento premium mensual', 'Hasta 4h de cambios/mejoras al mes, prioridad', '80–120 EUR/mes'],
            ['Dominio personalizado (.com/.es)', 'Gestion y configuracion del dominio', '10-15 EUR/ano + 30 EUR config'],
            ['Email profesional', 'info@negocio.com con Google Workspace o Zoho', '6-12 EUR/mes (Google)'],
            ['Google My Business', 'Creacion y optimizacion de la ficha de Google Maps', '60–100 EUR (una vez)'],
            ['SEO mensual basico', 'Optimizacion de contenido, palabras clave, informes mensuales', '100–200 EUR/mes'],
            ['Meta Ads (gestion campanas)', 'Gestion de campanas de publicidad en Instagram/Facebook', '100 EUR/mes + presupuesto'],
            ['Google Ads (gestion campanas)', 'Gestion de campanas en buscador', '100 EUR/mes + presupuesto'],
            ['Fotografia profesional del negocio', 'Sesion de fotos del local y del equipo', '150–300 EUR'],
            ['Pack redes sociales (1 mes)', 'Diseno de 12 posts + copies + calendario editorial', '200–350 EUR'],
        ]
    )
    divider(doc)

    # 3.4 PAGO Y CONDICIONES
    h2(doc, '3.4 Condiciones de pago habituales')
    highlight_box(doc, 'Politica estandar: 50% al inicio del proyecto + 50% al entregar la web funcionando.', 'E8F0FA', AZUL)

    table(doc,
        ['Condicion', 'Detalle'],
        [
            ['Forma de pago', 'Transferencia bancaria o Bizum'],
            ['Anticipo', '50% al confirmar el proyecto'],
            ['Pago final', '50% al entregar la web funcionando (antes del deploy final)'],
            ['Pago fraccionado', 'Disponible en 2 o 3 cuotas para proyectos > 1.000 EUR (sin intereses)'],
            ['Revision incluida', '2 rondas de revision incluidas en todos los paquetes'],
            ['Garantia post-entrega', '30 dias de soporte gratuito para bugs o errores del desarrollo'],
            ['Propiedad', 'El codigo es del cliente desde el primer dia'],
            ['Dominio y hosting', 'El cliente paga sus propias cuentas (Vercel es gratis, dominio ~12 EUR/ano)'],
        ]
    )
    divider(doc)

    # 3.5 CASO TWIN BROS (referencia)
    h2(doc, '3.5 Twin Bros como caso de referencia')
    para(doc, 'El proyecto Twin Bros es el mejor ejemplo del Pack Premium. Datos del proyecto real:', size=11)
    doc.add_paragraph()
    table(doc,
        ['Dato', 'Valor'],
        [
            ['Paquete equivalente', 'Pack Premium (web con sistema de citas + 4 artistas)'],
            ['Stack', 'Next.js 16 + React 19 + Tailwind v4 + Framer Motion'],
            ['Hosting', 'Vercel (gratuito) — https://twinbros-web.vercel.app'],
            ['Artistas incluidos', '4 (Pedro, Adri, Maranhas, Dual Piercing)'],
            ['Fotos reales', '17 fotos de Instagram'],
            ['Videos reales', '8 reels de Instagram (2 por artista)'],
            ['Panel admin', 'Si — citas, artistas, configuracion'],
            ['Sistema de citas', 'Si — formulario completo + panel de gestion'],
            ['Tiempo de desarrollo', '~5 dias (con descarga de media incluida)'],
            ['Estado', 'Entregado y en produccion desde 21/03/2026'],
        ]
    )
    divider(doc)
    highlight_box(doc, 'Twin Bros = referencia visual perfecta para ensenar a clientes del sector belleza, tatuaje, peluqueria, estetica o cualquier servicio con multiples profesionales.', 'F3E5FF', MAGENTA)

    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # SECCIÓN 4 — HOJA DE ARGUMENTARIO RAPIDO
    # ═════════════════════════════════════════════════════════════════════════
    h1(doc, '4. ARGUMENTARIO RAPIDO PARA LA VENTA WEB')
    para(doc, 'Para usar al presentar cualquier web de negocio local, adaptando al sector.', italic=True, size=10, color=GRIS)
    divider(doc)

    h2(doc, '4.1 Por que necesitan una web profesional')
    nbullet(doc, 'El 76% de los consumidores buscan en Google antes de ir a un negocio local.')
    nbullet(doc, 'Instagram no reemplaza la web: el algoritmo cambia, las cuentas pueden banearse, las busquedas de Google no indexan IG.')
    nbullet(doc, 'Una web con sistema de citas funciona 24h, los 7 dias, sin que el dueno tenga que estar pendiente del movil.')
    nbullet(doc, 'Cada semana sin web es trafico que se va a la competencia que SI tiene presencia online.')
    divider(doc)

    h2(doc, '4.2 Respuestas a objeciones frecuentes')
    objections = [
        ('"Ya tengo Instagram, para que quiero web"',
         'Instagram te da visibilidad entre tus seguidores. La web te da clientes nuevos que te buscan en Google. Son dos canales distintos. Ademas, si Instagram cambia el algoritmo o te suspende la cuenta, pierdes todo de golpe. Con la web, el cliente es tuyo.'),
        ('"Es muy caro para un negocio pequeno"',
         'Una sola cita nueva al mes generada por la web se paga sola en 6-12 meses. Twin Bros tiene ya la web funcionando y puede recibir citas online mientras duermen. La pregunta es cuanto les cuesta NO tenerla.'),
        ('"Wix o Squarespace me lo hace mas barato"',
         'Con Wix pagas todos los meses (15-25 EUR/mes = 180-300 EUR/ano) por un diseno generico que comparte con miles de negocios. Nosotros hacemos un diseno unico, tuyo para siempre, sin cuotas, en Vercel gratuito.'),
        ('"No tengo tiempo para gestionar una web"',
         'Por eso hacemos el admin tan sencillo. Actualizar un precio o anadir una foto tarda menos de 2 minutos. Y si no quieren tocarlo, tienen el plan de mantenimiento donde lo hacemos nosotros.'),
        ('"Necesito pensarlo"',
         'Perfecto. Te mando la propuesta por WhatsApp y te llamo el [dia concreto] para resolver dudas. Sin fecha de seguimiento no hay venta.'),
    ]
    for obj, resp in objections:
        para(doc, f'Objecion: {obj}', bold=True, size=10, color=ROJO)
        para(doc, f'Respuesta: {resp}', size=10, indent=0.3)
        doc.add_paragraph()

    h2(doc, '4.3 Checklist de cierre')
    checks = [
        'Has ensenado una web de referencia (Twin Bros, Chantarela, Bar Ryky)',
        'Has preguntado cuantos profesionales o servicios tiene el negocio',
        'Has identificado si necesitan citas online o solo presencia',
        'Has ofrecido el paquete correcto (Basico / Estandar / Premium)',
        'Has explicado que el hosting es gratis con Vercel',
        'Has mencionado que el codigo es suyo desde el primer dia',
        'Has ofrecido pago en 2 cuotas si hay resistencia al precio',
        'Has fijado una fecha de entrega concreta',
        'Has enviado la propuesta por escrito (WhatsApp o email)',
        'Tienes el nombre, telefono y email del responsable',
    ]
    for c in checks:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(f'[ ]  {c}')
        r.font.size = Pt(10.5)

    doc.add_paragraph()
    highlight_box(doc, 'Recordatorio: siempre mostrar primero una demo visual. Una web bien ensenada se vende sola.', '1A1A1A', BLANCO)

    doc.save(OUTPUT_PATH)
    print(f'Generado: {OUTPUT_PATH}')

if __name__ == '__main__':
    build()

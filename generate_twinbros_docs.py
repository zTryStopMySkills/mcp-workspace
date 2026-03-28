"""
Genera 3 documentos DOCX para Twin Bros Tattoo Studio
- DOC1: Presentación de la web (para que Luismi explique al cliente)
- DOC2: Mejoras posibles y hoja de ruta
- DOC3: Marketing digital, SEO y publicidad de pago
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = r"C:\Users\jose2\OneDrive\Escritorio\mcp"

# ─── HELPERS ─────────────────────────────────────────────────────────────────

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_para(doc, text, bold=False, italic=False, size=11, color=None, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(level * 0.5)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        # Black background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1A1A2E')
        tcPr.append(shd)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, cell_text in enumerate(row_data):
            row.cells[ci].text = str(cell_text)
    return table

def add_separator(doc):
    doc.add_paragraph('─' * 70)

def set_doc_margins(doc, top=2, bottom=2, left=2.5, right=2.5):
    section = doc.sections[0]
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)

# ─── DOC 1: PRESENTACIÓN DE LA WEB ───────────────────────────────────────────

def create_doc1():
    doc = Document()
    set_doc_margins(doc)

    # Portada
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('TWIN BROS TATTOO STUDIO')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(224, 64, 251)  # magenta neón

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run('Presentación de la Web Oficial')
    run2.font.size = Pt(16)
    run2.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('Documento preparado para explicar el proyecto al cliente\n').font.size = Pt(10)
    info.add_run('Tomares, Sevilla — Marzo 2026').font.size = Pt(10)

    doc.add_page_break()

    # ── 1. RESUMEN DEL PROYECTO ──
    add_heading(doc, '1. RESUMEN DEL PROYECTO', 1)
    add_para(doc, 'Twin Bros Tattoo Studio es un estudio de tatuaje y piercing ubicado en la Alameda Santa Eufemia N32, Tomares (Sevilla). El proyecto consiste en una web profesional completa que centraliza toda la presencia digital del estudio: 4 artistas con perfiles individuales, galería de trabajos, sistema de citas online y panel de administración.', size=11)
    doc.add_paragraph()

    add_heading(doc, '¿Qué es exactamente lo que se ha desarrollado?', 2)
    bullets = [
        'Una web profesional accesible desde cualquier dispositivo (móvil, tablet, ordenador)',
        'Diseño único con identidad visual propia del estudio: negro + magenta neón, fuente Bungee, estética urbana',
        'Perfiles individuales para los 4 artistas con galería de fotos y vídeos reales de Instagram',
        'Sistema de citas online: los clientes solicitan cita directamente desde la web',
        'Panel de administración para gestionar citas y contenido sin tocar código',
        'Disponible online desde el primer día, sin coste de hosting adicional (Vercel)',
    ]
    for b in bullets:
        add_bullet(doc, b)

    doc.add_paragraph()

    # ── 2. URL Y ACCESOS ──
    add_heading(doc, '2. URL Y ACCESOS', 1)
    add_table(doc,
        ['Elemento', 'Valor', 'Notas'],
        [
            ['URL de la web', 'https://twinbros-web.vercel.app', 'Accesible desde cualquier navegador'],
            ['Panel de administración', 'https://twinbros-web.vercel.app/admin', 'Solo para el equipo del estudio'],
            ['Contraseña admin', 'twinbros2026', 'Se puede cambiar cuando quieran'],
            ['Instagram estudio', '@twinbrostattoo', 'Enlazado desde la web'],
            ['Sponsor', '@bigwasp.official', 'Big Wasp, aparece en la web'],
        ]
    )
    doc.add_paragraph()
    add_para(doc, '⚠️  IMPORTANTE: El dominio actual es de demo (vercel.app). Para tener un dominio propio (ej: twinbrostattoo.com) se necesita contratarlo por separado — ver apartado de mejoras.', italic=True, size=10)

    doc.add_paragraph()

    # ── 3. SECCIONES DE LA WEB ──
    add_heading(doc, '3. SECCIONES DE LA WEB', 1)
    add_para(doc, 'La web está compuesta por las siguientes secciones, accesibles desde el menú de navegación:', size=11)
    doc.add_paragraph()

    sections_data = [
        ('HERO — Portada principal',
         'Lo primero que ve el visitante. Vídeo o imagen de fondo de gran impacto con el nombre del estudio, slogan y botón de reserva de cita. Los 4 artistas aparecen como accesos rápidos con sus colores identificativos.'),
        ('SOBRE NOSOTROS (About)',
         'Historia del estudio. Pedro y Adri como hermanos fundadores ("los Twin Bros"), filosofía del estudio, sponsor Big Wasp. Transmite la personalidad de marca.'),
        ('ARTISTAS',
         'Sección clave. Tarjetas interactivas para los 4 profesionales: Pedro Martín Art (Realismo), Adri Xegg (Anime & Color), Maranhas (Estilo internacional), Dual Piercing. Cada tarjeta muestra especialidades, bio y enlace a perfil individual.'),
        ('GALERÍA',
         'Filtro por artista. El visitante puede ver los trabajos de cada profesional por separado. Fotos y vídeos reales descargados de Instagram.'),
        ('CITAS / RESERVAS',
         'Formulario de solicitud de cita. El cliente elige artista, fecha, describe su idea, zona del cuerpo y tamaño. La cita llega directamente al panel de administración.'),
        ('CONTACTO',
         'Dirección del estudio con mapa interactivo (Google Maps), horario de apertura y enlaces a redes sociales.'),
    ]

    for title_s, desc in sections_data:
        add_para(doc, f'▸ {title_s}', bold=True, size=11)
        add_para(doc, desc, size=10, indent=0.5)
        doc.add_paragraph()

    # ── 4. PERFILES INDIVIDUALES DE ARTISTAS ──
    add_heading(doc, '4. PERFILES INDIVIDUALES DE ARTISTAS', 1)
    add_para(doc, 'Cada artista tiene su propia página dentro de la web (/artista/[nombre]), accesible desde la sección de artistas. Esto es importante para el posicionamiento en Google y para que cada artista pueda compartir su perfil en redes sociales.', size=11)
    doc.add_paragraph()

    artists = [
        ('Pedro Martín Art', '@pedromartinart', 'Realismo & Black and Grey', '16K seguidores en Instagram', 'Azul (#4FC3F7)'),
        ('Adri Xegg', '@adrixegg', 'Anime & Color', 'Hermano de Pedro, fundador', 'Naranja (#FF5722)'),
        ('Eduardo "Maranhas"', '@maranhastattoo', 'Estilo internacional', 'Visión cosmopolita', 'Verde (#66BB6A)'),
        ('Dual Piercing', '@dualpiercing', 'Piercings y joyería corporal', 'Especialista en perforación', 'Dorado (#FFD700)'),
    ]

    add_table(doc,
        ['Artista', 'Instagram', 'Especialidad', 'Detalle', 'Color'],
        artists
    )

    doc.add_paragraph()

    # ── 5. PANEL DE ADMINISTRACIÓN ──
    add_heading(doc, '5. PANEL DE ADMINISTRACIÓN', 1)
    add_para(doc, 'El panel de administración está disponible en /admin y es la herramienta de gestión del estudio. Desde aquí pueden hacer todo sin tocar código.', size=11)
    doc.add_paragraph()

    add_heading(doc, 'Dashboard — Inicio', 2)
    add_para(doc, 'Vista general con el número de citas pendientes, confirmadas y canceladas del día/semana.', size=10)

    add_heading(doc, 'Gestión de Citas', 2)
    bullets2 = [
        'Ver todas las solicitudes de cita con nombre del cliente, artista elegido, fecha y descripción',
        'Cambiar el estado de cada cita: Pendiente → Confirmada / Cancelada',
        'Eliminar citas antiguas',
        'Filtrar por artista o por estado',
    ]
    for b in bullets2:
        add_bullet(doc, b)

    add_heading(doc, 'Gestión de Artistas', 2)
    bullets3 = [
        'Editar la bio de cada artista',
        'Cambiar la foto de perfil',
        'Actualizar sus especialidades',
        'Gestionar sus enlaces de redes sociales',
    ]
    for b in bullets3:
        add_bullet(doc, b)

    add_heading(doc, 'Configuración del estudio', 2)
    bullets4 = [
        'Cambiar dirección y horarios',
        'Actualizar datos de contacto',
        'Gestionar el sponsor (Big Wasp)',
    ]
    for b in bullets4:
        add_bullet(doc, b)

    doc.add_paragraph()
    add_para(doc, 'El panel está protegido por contraseña. Solo los miembros del estudio tienen acceso. Los cambios se aplican en tiempo real a la web sin necesidad de republicar nada.', italic=True, size=10)

    doc.add_paragraph()

    # ── 6. TECNOLOGÍA ──
    add_heading(doc, '6. TECNOLOGÍA UTILIZADA', 1)
    add_para(doc, 'La web está construida con tecnología de nivel profesional, la misma que usan grandes empresas tecnológicas:', size=11)
    doc.add_paragraph()

    add_table(doc,
        ['Tecnología', 'Para qué sirve'],
        [
            ['Next.js 16 + React 19', 'Framework web moderno de alto rendimiento (usado por Nike, TikTok, Twitch)'],
            ['Tailwind CSS v4', 'Diseño responsive y consistente en todos los dispositivos'],
            ['Framer Motion', 'Animaciones fluidas y profesionales'],
            ['Vercel', 'Hosting gratuito con CDN global, carga ultrarrápida'],
            ['TypeScript', 'Código robusto y sin errores inesperados'],
        ]
    )

    doc.add_paragraph()
    add_para(doc, 'El hosting en Vercel (plan gratuito) es suficiente para un estudio de tatuaje. El sitio carga en menos de 2 segundos en todo el mundo.', size=10)

    doc.add_paragraph()

    # ── 7. CONTENIDO MULTIMEDIA ──
    add_heading(doc, '7. CONTENIDO MULTIMEDIA', 1)
    add_para(doc, 'La web utiliza contenido real del estudio, no imágenes genéricas de stock:', size=11)
    doc.add_paragraph()
    add_table(doc,
        ['Tipo', 'Cantidad', 'Procedencia'],
        [
            ['Fotos de trabajos', '17 fotografías reales', 'Descargadas de Instagram'],
            ['Vídeos reels', '8 vídeos (2 por artista)', 'Reels de Instagram'],
            ['Fotos de convenciones', '4 (Sevilla + Jerez)', 'Instagram estudio'],
            ['Foto sponsor', '1 (Big Wasp)', 'Instagram @bigwasp.official'],
            ['Logo estudio', 'Monograma gótico TB', 'Identidad visual del estudio'],
        ]
    )

    doc.add_page_break()

    # ── 8. RESUMEN PARA EL CLIENTE ──
    add_heading(doc, '8. RESUMEN — LO QUE TIENE EL ESTUDIO AHORA', 1)
    add_para(doc, 'En una sola frase: Twin Bros Tattoo tiene ahora una web profesional que trabaja para ellos 24/7, capta clientes nuevos, gestiona sus citas y representa el nivel real del estudio.', bold=True, size=12)
    doc.add_paragraph()

    add_table(doc,
        ['Qué tiene', 'Estado'],
        [
            ['Web online y accesible', '✅ Funcionando'],
            ['Diseño único y profesional', '✅ Completado'],
            ['4 perfiles de artistas con galería', '✅ Completado'],
            ['Fotos y vídeos reales de Instagram', '✅ Completado'],
            ['Sistema de citas online', '✅ Completado'],
            ['Panel de administración', '✅ Completado'],
            ['Optimización para móvil', '✅ Completado'],
            ['Hosting gratuito activo', '✅ Vercel'],
            ['Dominio personalizado (.com)', '⚠️ Pendiente (ver mejoras)'],
            ['Posicionamiento en Google', '⚠️ Pendiente (ver marketing)'],
        ]
    )

    path = os.path.join(OUTPUT_DIR, 'TwinBros_01_Presentacion_Web.docx')
    doc.save(path)
    print(f'✅ DOC 1 generado: {path}')
    return path


# ─── DOC 2: MEJORAS POSIBLES ─────────────────────────────────────────────────

def create_doc2():
    doc = Document()
    set_doc_margins(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('TWIN BROS TATTOO STUDIO')
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(224, 64, 251)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run('Mejoras Posibles y Hoja de Ruta')
    run2.font.size = Pt(16)

    doc.add_paragraph()
    doc.add_page_break()

    # Intro
    add_heading(doc, 'INTRODUCCIÓN', 1)
    add_para(doc, 'La web actual es completamente funcional y está lista para usar. Las siguientes mejoras son opcionales pero recomendadas para maximizar el impacto del estudio online. Se clasifican por prioridad y coste estimado.', size=11)
    doc.add_paragraph()

    # ── CATEGORÍA 1: PRIORITARIO ──
    add_heading(doc, '1. MEJORAS PRIORITARIAS (Alto impacto, bajo coste)', 1)
    doc.add_paragraph()

    add_heading(doc, '1.1 Dominio personalizado', 2)
    add_para(doc, 'Estado actual: la web está en twinbros-web.vercel.app (dominio gratuito de demo)', size=10, italic=True)
    add_para(doc, 'Mejora: Contratar el dominio twinbrostattoo.com o twinbros.es para:', size=11)
    bullets = [
        'Aspecto 100% profesional y memorable',
        'Mayor confianza de los clientes',
        'Mejor posicionamiento en Google',
        'Email profesional (info@twinbrostattoo.com)',
    ]
    for b in bullets:
        add_bullet(doc, b)
    add_para(doc, '💶 Coste: 10–15€/año (GoDaddy, Namecheap, Nominalia)', bold=True, size=10)
    add_para(doc, '⏱ Tiempo de implementación: 1–2 horas (configuración DNS)', size=10, italic=True)
    doc.add_paragraph()

    add_heading(doc, '1.2 Google My Business (ficha de empresa gratuita)', 2)
    add_para(doc, 'Crear y verificar la ficha de Twin Bros en Google Maps. Esto hace que aparezcan en el mapa cuando alguien busca "tatuajes Tomares" o "tattoo Sevilla".', size=11)
    bullets2 = [
        'Aparición en Google Maps con fotos, horario y valoraciones',
        'Los clientes pueden dejar reseñas directamente en Google',
        'Enlace directo a la web desde la ficha',
        'Aparición en el bloque de resultados locales (muy visible)',
    ]
    for b in bullets2:
        add_bullet(doc, b)
    add_para(doc, '💶 Coste: GRATIS (requiere verificación por tarjeta postal)', bold=True, size=10)
    add_para(doc, '⏱ Tiempo: 1–2 semanas (verificación por correo postal de Google)', size=10, italic=True)
    doc.add_paragraph()

    add_heading(doc, '1.3 WhatsApp Business integrado en la web', 2)
    add_para(doc, 'Actualmente la web no tiene WhatsApp. Muchos clientes prefieren contactar por WhatsApp antes de rellenar un formulario.', size=11)
    add_para(doc, 'Mejora: Añadir botón flotante de WhatsApp Business que abra conversación directa con el número del estudio.', size=11)
    add_para(doc, '💶 Coste: 0€ (solo desarrollo — estimado 1 hora de trabajo)', bold=True, size=10)
    add_para(doc, '⏱ Tiempo de implementación: 1 hora', size=10, italic=True)
    doc.add_paragraph()

    add_heading(doc, '1.4 Formulario de citas con notificación por email', 2)
    add_para(doc, 'Actualmente las citas se guardan en el panel de admin pero no envían un email automático al estudio cuando llega una nueva solicitud.', size=11)
    add_para(doc, 'Mejora: Envío automático de email al estudio cuando hay cita nueva + email de confirmación al cliente.', size=11)
    add_para(doc, '💶 Coste: 0€ (usando Resend o EmailJS, plan gratuito)', bold=True, size=10)
    add_para(doc, '⏱ Tiempo de implementación: 2–3 horas', size=10, italic=True)
    doc.add_paragraph()

    # ── CATEGORÍA 2: MEDIA PRIORIDAD ──
    add_heading(doc, '2. MEJORAS DE MEDIA PRIORIDAD (Impacto medio-alto)', 1)
    doc.add_paragraph()

    add_heading(doc, '2.1 Blog / Artículo de contenido', 2)
    add_para(doc, 'Añadir una sección de blog con artículos sobre tatuajes: "Cómo cuidar un tatuaje nuevo", "Estilos de tatuaje explicados", "El antes y después de Pedro Martín". Este contenido posiciona muy bien en Google.', size=11)
    add_para(doc, '💶 Coste de desarrollo: 200–400€ (una sola vez)', bold=True, size=10)
    add_para(doc, 'Redacción de artículos: puede hacerse con IA (ChatGPT) o contratar redactor freelance (20–50€ por artículo)', size=10, italic=True)
    doc.add_paragraph()

    add_heading(doc, '2.2 Galería mejorada con lightbox y filtros avanzados', 2)
    add_para(doc, 'La galería actual muestra fotos por artista. Se puede mejorar con:', size=11)
    bullets3 = [
        'Lightbox (foto en pantalla completa al hacer clic)',
        'Filtros por estilo: Realismo, Anime, Color, Black & Grey, Piercing',
        'Función de "antes y después" para mostrar evolución de un tatuaje',
        '"Te puede interesar" — sugerencias de trabajos similares',
    ]
    for b in bullets3:
        add_bullet(doc, b)
    add_para(doc, '💶 Coste estimado: 150–300€', bold=True, size=10)
    doc.add_paragraph()

    add_heading(doc, '2.3 Sistema de citas avanzado', 2)
    add_para(doc, 'El sistema actual recibe solicitudes. Una mejora importante sería:', size=11)
    bullets4 = [
        'Calendario visual de disponibilidad por artista (el cliente ve huecos libres)',
        'Depósito online para confirmar cita (Stripe/PayPal — 10–20€ señal)',
        'Recordatorio automático por SMS/WhatsApp 24h antes de la cita',
        'Historial de clientes (quién ha venido, cuántas veces, qué trabajos)',
    ]
    for b in bullets4:
        add_bullet(doc, b)
    add_para(doc, '💶 Coste estimado: 400–800€ (desarrollo más complejo)', bold=True, size=10)
    doc.add_paragraph()

    add_heading(doc, '2.4 Tienda online / Merchandising', 2)
    add_para(doc, 'Si el estudio vende productos (aftercare, merchandise, prints de los artistas), se puede añadir una tienda online básica.', size=11)
    add_para(doc, '💶 Coste estimado: 300–600€ (con Stripe integrado)', bold=True, size=10)
    doc.add_paragraph()

    # ── CATEGORÍA 3: LARGO PLAZO ──
    add_heading(doc, '3. MEJORAS A LARGO PLAZO', 1)
    doc.add_paragraph()

    add_heading(doc, '3.1 App móvil nativa', 2)
    add_para(doc, 'Una app para iOS y Android con notificaciones push para recordatorios de cita, seguimiento del proceso de cicatrización, galería privada del cliente.', size=11)
    add_para(doc, '💶 Coste estimado: 3.000–8.000€ (desarrollo externo) | 1.500–3.000€ (PWA avanzada)', bold=True, size=10)
    doc.add_paragraph()

    add_heading(doc, '3.2 Plataforma multi-estudio', 2)
    add_para(doc, 'Si el negocio crece y abren más locales o se asocian con otros estudios, la plataforma se puede escalar a una red de estudios bajo Twin Bros.', size=11)
    add_para(doc, '💶 Coste estimado: 5.000–15.000€', bold=True, size=10)
    doc.add_paragraph()

    # ── RESUMEN DE INVERSIÓN ──
    add_heading(doc, '4. RESUMEN DE INVERSIÓN POR MEJORAS', 1)
    add_table(doc,
        ['Mejora', 'Prioridad', 'Coste', 'ROI esperado'],
        [
            ['Dominio propio (.com)', 'ALTA', '10-15€/año', 'Credibilidad inmediata'],
            ['Google My Business', 'ALTA', 'GRATIS', 'Clientes locales nuevos'],
            ['WhatsApp en la web', 'ALTA', '0€', 'Más conversiones'],
            ['Notificaciones email citas', 'ALTA', '0€', 'No perder citas'],
            ['Blog de contenido', 'MEDIA', '200-400€', 'SEO orgánico a 6 meses'],
            ['Galería mejorada', 'MEDIA', '150-300€', 'Más tiempo en web'],
            ['Sistema citas avanzado', 'MEDIA', '400-800€', 'Menos no-shows'],
            ['Tienda online', 'BAJA', '300-600€', 'Ingresos adicionales'],
            ['App móvil', 'LARGO PLAZO', '1500-8000€', 'Retención de clientes'],
        ]
    )

    path = os.path.join(OUTPUT_DIR, 'TwinBros_02_Mejoras_y_Hoja_de_Ruta.docx')
    doc.save(path)
    print(f'✅ DOC 2 generado: {path}')
    return path


# ─── DOC 3: MARKETING DIGITAL ────────────────────────────────────────────────

def create_doc3():
    doc = Document()
    set_doc_margins(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('TWIN BROS TATTOO STUDIO')
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(224, 64, 251)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run('Estrategia de Marketing Digital\nSEO, Google Ads y Publicidad de Pago')
    run2.font.size = Pt(16)

    doc.add_paragraph()
    doc.add_page_break()

    # INTRO
    add_heading(doc, 'SITUACIÓN ACTUAL', 1)
    add_para(doc, 'Twin Bros Tattoo tiene ya la web online y un perfil de Instagram activo (@twinbrostattoo). El siguiente paso es hacer que los clientes potenciales encuentren el estudio cuando buscan en Google o en Instagram.', size=11)
    doc.add_paragraph()
    add_para(doc, 'Hay dos vías principales:', bold=True, size=11)
    add_bullet(doc, 'SEO (posicionamiento orgánico): aparecer en Google de forma gratuita pero a largo plazo')
    add_bullet(doc, 'Publicidad de pago (SEM / Social Ads): aparecer arriba inmediatamente, pagando')
    doc.add_paragraph()

    # ── 1. SEO LOCAL ──
    add_heading(doc, '1. SEO LOCAL — POSICIONAMIENTO EN GOOGLE', 1)
    add_para(doc, 'El SEO local es la estrategia para que Twin Bros aparezca cuando alguien en Sevilla/Tomares busca servicios de tatuaje. Es la inversión con mejor retorno a largo plazo.', size=11)
    doc.add_paragraph()

    add_heading(doc, '1.1 Palabras clave objetivo', 2)
    add_para(doc, 'Estas son las búsquedas que hacen los clientes potenciales:', size=10)
    add_table(doc,
        ['Palabra clave', 'Búsquedas/mes (estimadas)', 'Dificultad', 'Prioridad'],
        [
            ['"tatuajes Sevilla"', '1.000–3.000', 'Alta', 'Principal'],
            ['"tatuajes Tomares"', '100–300', 'Baja', 'Principal'],
            ['"tatuaje realismo Sevilla"', '200–500', 'Media', 'Principal'],
            ['"tatuaje anime Sevilla"', '100–300', 'Baja', 'Principal'],
            ['"tatuador black and grey Sevilla"', '50–150', 'Baja', 'Secundaria'],
            ['"estudios de tatuaje cerca de mí"', '500–1.500', 'Alta', 'Oportunidad'],
            ['"piercing Tomares"', '50–200', 'Muy baja', 'Fácil ganar'],
            ['"piercing Sevilla"', '200–500', 'Media', 'Secundaria'],
        ]
    )
    doc.add_paragraph()

    add_heading(doc, '1.2 Acciones SEO recomendadas', 2)

    add_para(doc, 'A. Google My Business (GRATIS — PRIORITARIO)', bold=True, size=11)
    add_para(doc, 'Esta es la acción más importante y es gratuita. Al crear la ficha de empresa en Google:', size=10)
    bullets = [
        'Twin Bros aparece en Google Maps cuando alguien busca "tatuajes cerca"',
        'Los clientes pueden ver fotos, horarios, reseñas y llamar directamente desde el buscador',
        'Es el factor nº1 de posicionamiento local',
        'Coste: GRATIS. Solo requiere un par de horas de configuración y verificación postal (2 semanas)',
    ]
    for b in bullets:
        add_bullet(doc, b)
    doc.add_paragraph()

    add_para(doc, 'B. Optimización SEO On-Page de la web (ya incluido en la web actual)', bold=True, size=11)
    add_para(doc, 'La web ya incluye etiquetas meta, estructura semántica y textos optimizados para posicionamiento. Esto es la base.', size=10)
    doc.add_paragraph()

    add_para(doc, 'C. Blog de contenido (inversión recomendada)', bold=True, size=11)
    add_para(doc, 'Artículos de blog posicionan en Google con búsquedas específicas de alto valor:', size=10)
    topics = [
        '"Cómo cuidar un tatuaje nuevo" — miles de búsquedas mensuales',
        '"Diferencia entre tatuaje realismo y black & grey"',
        '"Cuánto cuesta un tatuaje en Sevilla"',
        '"Los mejores tatuadores de Sevilla"',
        '"Tatuaje de anime: qué estilo elegir"',
    ]
    for t in topics:
        add_bullet(doc, t)
    add_para(doc, 'Cada artículo es tráfico orgánico gratuito de por vida.', italic=True, size=10)
    doc.add_paragraph()

    add_heading(doc, '1.3 Coste del SEO y plazos realistas', 2)
    add_table(doc,
        ['Acción', 'Coste', 'Tiempo para ver resultados'],
        [
            ['Google My Business (config + gestión mensual)', '0€ / 50–100€/mes si se externaliza', '2–4 semanas'],
            ['Optimización web básica', 'Ya incluida en la web', 'Inmediato'],
            ['Blog (5 artículos iniciales)', '0€ con IA / 150–250€ con redactor', '3–6 meses'],
            ['Link building (otros webs enlacen)', '0€ orgánico / 100–300€/mes agencia', '6–12 meses'],
            ['Servicio SEO profesional completo', '200–600€/mes (agencia)', '4–8 meses'],
        ]
    )
    doc.add_paragraph()
    add_para(doc, 'REALIDAD DEL SEO: El SEO no da resultados en semanas. El posicionamiento orgánico tarda entre 3 y 9 meses en dar frutos notables. Es una inversión a largo plazo con retorno indefinido.', bold=True, size=10)

    doc.add_paragraph()

    # ── 2. GOOGLE ADS ──
    add_heading(doc, '2. GOOGLE ADS — PUBLICIDAD EN BUSCADOR', 1)
    add_para(doc, 'Google Ads permite aparecer en la primera posición de Google inmediatamente, pagando por cada clic. Es la opción para obtener clientes en semanas, no meses.', size=11)
    doc.add_paragraph()

    add_heading(doc, '2.1 Cómo funciona', 2)
    add_para(doc, 'Cuando alguien escribe "tatuajes Sevilla" en Google, los primeros resultados (marcados con "Patrocinado") son anuncios de pago. Twin Bros puede aparecer ahí pagando por clic.', size=11)
    doc.add_paragraph()

    add_heading(doc, '2.2 Coste estimado por clic (CPC) en el sector tatuaje', 2)
    add_table(doc,
        ['Palabra clave', 'CPC estimado', 'Competencia'],
        [
            ['"tatuajes Sevilla"', '0,50–1,50€ por clic', 'Media-alta'],
            ['"tatuaje realismo Sevilla"', '0,30–0,80€ por clic', 'Baja'],
            ['"tatuaje anime Sevilla"', '0,20–0,60€ por clic', 'Muy baja'],
            ['"piercing Sevilla"', '0,20–0,50€ por clic', 'Baja'],
            ['"tatuaje Tomares"', '0,10–0,40€ por clic', 'Muy baja'],
        ]
    )
    doc.add_paragraph()

    add_heading(doc, '2.3 Presupuestos recomendados', 2)
    add_table(doc,
        ['Plan', 'Presupuesto mensual', 'Clics estimados/mes', 'Clientes potenciales estimados'],
        [
            ['Básico (prueba)', '50–100€/mes', '60–150 clics', '5–15 consultas'],
            ['Estándar', '150–300€/mes', '200–500 clics', '15–40 consultas'],
            ['Agresivo', '500–800€/mes', '700–1.500 clics', '50–100 consultas'],
        ]
    )
    doc.add_paragraph()
    add_para(doc, 'Estimación de conversión: en el sector tatuaje, aproximadamente el 5–10% de los que entran a la web acaban solicitando cita o consultando. Un tatuaje medio en Sevilla cuesta 100–400€. Con 15 consultas y una tasa de cierre del 30%, son 4–5 clientes nuevos al mes por 100€ de inversión.', size=10, italic=True)
    doc.add_paragraph()

    add_heading(doc, '2.4 Recomendación para empezar', 2)
    add_bullet(doc, 'Comenzar con 100€/mes durante 2 meses como prueba')
    add_bullet(doc, 'Segmentar por ubicación: Tomares, Sevilla capital, municipios limítrofes')
    add_bullet(doc, 'Anuncios centrados en los puntos fuertes: Pedro Martín realismo, Adri Xegg anime')
    add_bullet(doc, 'Optimizar el anuncio cada 2 semanas según resultados')
    doc.add_paragraph()
    add_para(doc, '⚠️ Importante: Google Ads requiere seguimiento constante. Sin optimización, el dinero se gasta sin retorno. Se recomienda gestión por alguien que entienda la plataforma o contratar agencia para presupuestos superiores a 300€/mes.', italic=True, size=10)

    doc.add_paragraph()

    # ── 3. META ADS ──
    add_heading(doc, '3. META ADS — PUBLICIDAD EN INSTAGRAM Y FACEBOOK', 1)
    add_para(doc, 'Dado que los artistas ya tienen presencia en Instagram, Meta Ads (Instagram + Facebook) es la plataforma más natural y efectiva para Twin Bros. El contenido visual de tatuajes funciona especialmente bien en Instagram.', size=11)
    doc.add_paragraph()

    add_heading(doc, '3.1 Por qué Instagram es ideal para tatuajes', 2)
    bullets5 = [
        'El tatuaje es un producto 100% visual — Instagram es la red visual por excelencia',
        'Pedro Martín ya tiene 16.000 seguidores en IG (@pedromartinart) — base de audiencia ya creada',
        'Los anuncios de Instagram se integran naturalmente entre fotos de tatuajes',
        'Segmentación muy precisa: edad, intereses (tattoo, arte, manga, anime), ubicación',
        'Formato Reels patrocinado: enormemente efectivo con los vídeos que ya tienen del estudio',
    ]
    for b in bullets5:
        add_bullet(doc, b)
    doc.add_paragraph()

    add_heading(doc, '3.2 Tipos de campaña recomendados', 2)
    add_table(doc,
        ['Tipo de campaña', 'Objetivo', 'Formato', 'Coste estimado/mes'],
        [
            ['Reconocimiento de marca', 'Que la gente conozca el estudio', 'Imágenes + Reels', '50–100€'],
            ['Tráfico a la web', 'Llevar visitantes a la web', 'Carrusel de trabajos', '80–200€'],
            ['Generación de leads', 'Solicitudes de cita directas', 'Formulario en IG', '100–250€'],
            ['Retargeting', 'Re-impactar visitantes de la web', 'Recordatorio visual', '30–80€'],
        ]
    )
    doc.add_paragraph()

    add_heading(doc, '3.3 Audiencias clave para segmentar', 2)
    add_table(doc,
        ['Audiencia', 'Descripción'],
        [
            ['Jóvenes 18–35 en Sevilla y alrededores', 'Perfil principal del cliente de tatuaje'],
            ['Intereses: tattoo, arte, manga, anime, piercing', 'Alta intención de compra'],
            ['Lookalike de seguidores de @pedromartinart', 'Similar a quien ya les sigue'],
            ['Visitantes de la web (retargeting)', 'Ya mostraron interés — los más valiosos'],
            ['Fans de series anime (Naruto, One Piece, Dragon Ball)', 'Objetivo específico para Adri Xegg'],
        ]
    )
    doc.add_paragraph()

    add_heading(doc, '3.4 Presupuestos recomendados Meta Ads', 2)
    add_table(doc,
        ['Plan', 'Inversión mensual', 'Alcance estimado', 'Objetivo'],
        [
            ['Inicio', '50–100€/mes', '5.000–15.000 personas', 'Visibilidad local'],
            ['Crecimiento', '150–300€/mes', '20.000–50.000 personas', 'Citas + seguidores'],
            ['Escala', '400–700€/mes', '70.000–150.000 personas', 'Liderazgo de mercado'],
        ]
    )
    doc.add_paragraph()

    # ── 4. TIKTOK ADS ──
    add_heading(doc, '4. TIKTOK ADS — OPCIÓN ADICIONAL', 1)
    add_para(doc, 'Pedro Martín (@pedromartinart) y Adri Xegg (@adrixegg) ya tienen perfil en TikTok. Si generan contenido de proceso de tatuaje (timelapse, before/after), TikTok puede ser un canal muy potente.', size=11)
    doc.add_paragraph()

    add_para(doc, 'TikTok Ads funciona especialmente bien para:', size=11)
    add_bullet(doc, 'Vídeos de proceso de tatuaje (el contenido más viral en la plataforma)')
    add_bullet(doc, 'Público joven 16–28 años, muy afín al tatuaje')
    add_bullet(doc, 'Coste por mil impresiones más barato que Instagram en 2025–2026')
    doc.add_paragraph()
    add_para(doc, '💶 Inversión mínima recomendada: 50€/mes. Requiere crear contenido vídeo específico para TikTok.', bold=True, size=10)

    doc.add_paragraph()

    # ── 5. PLAN RECOMENDADO ──
    add_heading(doc, '5. PLAN DE MARKETING RECOMENDADO POR FASES', 1)
    doc.add_paragraph()

    add_heading(doc, 'FASE 1 — Fundamentos (Mes 1–2) — Coste: ~0€', 2)
    bullets6 = [
        'Crear y verificar Google My Business (GRATIS)',
        'Optimizar el perfil de Instagram existente con enlace a la web',
        'Añadir la web al pie de los perfiles de artistas (@pedromartinart, @adrixegg)',
        'Responder a todos los comentarios y DMs de Instagram',
        'Pedir a clientes satisfechos que dejen reseña en Google',
    ]
    for b in bullets6:
        add_bullet(doc, b)
    doc.add_paragraph()

    add_heading(doc, 'FASE 2 — Primera inversión (Mes 3–4) — Coste: 100–200€/mes', 2)
    bullets7 = [
        'Lanzar campaña Meta Ads de tráfico a la web (100€/mes)',
        'Boost de posts de Instagram con los mejores trabajos',
        'Comenzar a publicar 1 artículo de blog al mes en la web',
        'Seguimiento mensual de métricas (cuántas visitas, cuántas citas)',
    ]
    for b in bullets7:
        add_bullet(doc, b)
    doc.add_paragraph()

    add_heading(doc, 'FASE 3 — Escala (Mes 5+) — Coste: 300–600€/mes', 2)
    bullets8 = [
        'Añadir Google Ads (100–200€/mes)',
        'Aumentar Meta Ads con retargeting',
        'Blog activo (2–3 artículos/mes)',
        'Evaluar TikTok Ads si hay contenido de vídeo disponible',
        'Seguimiento semanal y optimización de campañas',
    ]
    for b in bullets8:
        add_bullet(doc, b)
    doc.add_paragraph()

    # ── RESUMEN FINAL ──
    add_heading(doc, '6. RESUMEN DE INVERSIÓN EN MARKETING', 1)
    add_table(doc,
        ['Canal', 'Coste mensual', 'Cuándo empezar', 'Retorno esperado'],
        [
            ['Google My Business', 'GRATIS', 'Inmediato', 'Clientes locales sin inversión'],
            ['SEO orgánico (blog)', '0–100€', 'Mes 1', 'Tráfico gratuito desde mes 4–6'],
            ['Meta Ads (Instagram)', '100–300€', 'Mes 2–3', '5–20 nuevas consultas/mes'],
            ['Google Ads', '100–300€', 'Mes 3–4', '5–20 nuevas consultas/mes'],
            ['TikTok Ads', '50–150€', 'Opcional', 'Alcance masivo jóvenes'],
            ['TOTAL inversión inicial', '100–200€/mes', 'Mes 1', 'ROI positivo estimado en 3 meses'],
            ['TOTAL escala', '400–800€/mes', 'Mes 5+', 'Liderazgo de mercado local'],
        ]
    )
    doc.add_paragraph()

    add_heading(doc, 'CONCLUSIÓN', 1)
    add_para(doc, 'Con una inversión inicial de 100–200€/mes en Meta Ads + Google My Business gratuito, Twin Bros puede generar 10–25 consultas nuevas al mes de clientes que no los conocían. Dado que el ticket medio de un tatuaje en Sevilla es 150–400€, el retorno de la inversión es rápido y claro.', bold=True, size=11)
    doc.add_paragraph()
    add_para(doc, 'La combinación ganadora: web profesional ✓ + Google My Business ✓ + Meta Ads activo = posicionamiento dominante en el mercado local de Tomares y Sevilla.', size=11, italic=True)

    path = os.path.join(OUTPUT_DIR, 'TwinBros_03_Marketing_Digital.docx')
    doc.save(path)
    print(f'✅ DOC 3 generado: {path}')
    return path


# ─── MAIN ─────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    print('Generando documentos Twin Bros Tattoo Studio...\n')
    p1 = create_doc1()
    p2 = create_doc2()
    p3 = create_doc3()
    print(f'\n✅ Todos los documentos generados en: {OUTPUT_DIR}')
    print(f'  - {os.path.basename(p1)}')
    print(f'  - {os.path.basename(p2)}')
    print(f'  - {os.path.basename(p3)}')
